import os
import hashlib
import re
from abc import ABC, abstractmethod
# 文档加载器
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader
)
from typing import List, Any, Optional
import xml.etree.ElementTree as ET
from langchain.docstore.document import Document
import json
import mammoth
from docx import Document as DocxDocument
# ================================
# 3. 数据处理抽象基类
# ================================
class DataProcessor(ABC):
    """数据处理器抽象基类，支持扩展"""

    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理该文件类型"""
        pass

    @abstractmethod
    def load_documents(self, file_path: str) -> List[Any]:
        """加载文档"""
        pass

    @abstractmethod
    def get_file_hash(self, file_path: str) -> str:
        """获取文件哈希值用于增量更新"""
        pass

    @property
    def custom_chunking(self) -> bool:
        """是否已实现自定义分块"""
        return False  # 默认未实现
# ================================
# 4. 数据处理工厂类
# ================================
class DataProcessorFactory:
    """数据处理器工厂类"""

    def __init__(self):
        self.processors = [
            PDFProcessor(),
            TextProcessor(),
            WordProcessor(),
            XMLProcessor(),
            JSONProcessor()
        ]

    def get_processor(self, file_path: str) -> Optional[DataProcessor]:
        """根据文件类型获取对应的处理器"""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
# ================================
# 5. 具体数据处理器实现
# ================================
class PDFProcessor(DataProcessor):
    """PDF文档处理器"""

    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理PDF文件"""
        if file_path.lower().endswith('.pdf'):
            print(f"检测到PDF文件: {os.path.basename(file_path)}")
            return True
        else:
            return False

    def load_documents(self, file_path: str) -> List[Any]:
        """加载并解析PDF文档内容"""
        print(f"开始处理PDF文档: {os.path.basename(file_path)}")
        try:
            loader = PyPDFLoader(file_path)
            raw_docs = loader.load()  # List[Document]，每页一个

            section_pattern = re.compile(r"^\s*\d+(\.\d+)*\s+[A-Z].+")
            sections = []
            current_title = None
            current_content = []

            for doc in raw_docs:
                lines = doc.page_content.splitlines()
                for line in lines:
                    if section_pattern.match(line.strip()):
                        # 发现新章节，保存旧章节
                        if current_title and current_content:
                            page_content = "\n".join(current_content).strip()
                            metadata = {"title": current_title.strip()}
                            sections.append(Document(page_content=page_content, metadata=metadata))
                            current_content = []
                        current_title = line.strip()
                    elif current_title:
                        current_content.append(line)

            # 最后一节
            if current_title and current_content:
                page_content = "\n".join(current_content).strip()
                metadata = {"title": current_title.strip()}
                sections.append(Document(page_content=page_content, metadata=metadata))

            return sections
        except Exception as e:
            print(f"处理失败：{e}")
            raise

    def get_file_hash(self, file_path: str) -> str:
        print(f"开始计算PDF文档hash: {os.path.basename(file_path)}")
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()

    @property
    def custom_chunking(self) -> bool:
        """已实现自定义分块"""
        return True

class TextProcessor(DataProcessor):
    """文本文档处理器"""
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理文本文件"""
        if file_path.lower().endswith(('.txt', '.md')):
            print(f"检测到文本文件: {os.path.basename(file_path)}")
            return True
        else:
            return False

    def load_documents(self, file_path: str) -> List[Any]:
        print(f"开始处理文本文档: {os.path.basename(file_path)}")
        loader = TextLoader(file_path, encoding='utf-8')
        text= loader.load()
        return text

    def get_file_hash(self, file_path: str) -> str:
        print(f"开始计算文本文档hash: {os.path.basename(file_path)}")
        with open(file_path, 'r', encoding='utf-8') as f:
            return hashlib.md5(f.read().encode()).hexdigest()

class WordProcessor(DataProcessor):
    """Word文档处理器"""
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理docx文件"""
        if file_path.lower().endswith('.docx'):
            print(f"检测到docx文件: {os.path.basename(file_path)}")
            return True
        else:
            return False

    def load_documents(self, file_path: str) -> List[Document]:
        """加载并解析Word文档内容，根据标题划分章节"""
        print(f"开始处理Word文档: {os.path.basename(file_path)}")

        try:
            # 验证文件路径
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            if not os.path.isfile(file_path):
                raise ValueError(f"路径不是文件: {file_path}")

            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                raise ValueError(f"文件为空: {file_path}")

            print(f"文件大小: {file_size} bytes")

            # 方法1：使用python-docx读取Word文档
            content = self._extract_content_with_docx(file_path)

            # 如果docx方法失败，尝试使用mammoth
            if not content.strip():
                print("docx方法未获取到内容，尝试使用mammoth...")
                content = self._extract_content_with_mammoth(file_path)

            if not content.strip():
                raise ValueError("无法从文档中提取任何内容")

            print(f"成功提取内容，长度: {len(content)} 字符")

            # 根据标题模式划分章节
            sections = self._split_by_headers(content, file_path)

            print(f"成功提取 {len(sections)} 个章节")
            return sections

        except Exception as e:
            print(f"处理失败：{e}")
            raise

    def _extract_content_with_docx(self, file_path: str) -> str:
        """使用python-docx提取Word文档内容"""
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 检查文件扩展名
            if not file_path.lower().endswith(('.docx', '.doc')):
                raise ValueError(f"不支持的文件格式: {file_path}")

            doc = DocxDocument(file_path)
            content_lines = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_lines.append(text)

            return "\n".join(content_lines)
        except Exception as e:
            print(f"使用docx提取失败: {e}")
            return ""

    def _extract_content_with_mammoth(self, file_path: str) -> str:
        """使用mammoth提取Word文档内容"""
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        except Exception as e:
            print(f"使用mammoth提取失败: {e}")
            return ""

    def _split_by_headers(self, content: str, file_path: str = None) -> List[Document]:
        """根据标题模式划分文档内容"""
        sections = []

        # 定义多种标题模式
        header_patterns = [
            re.compile(r"^\s*#{1,6}\s+.*$", re.MULTILINE),  # Markdown标题 (# ## ### 等)
            re.compile(r"^\s*\d+(\.\d+)*\s+[^\s].*$", re.MULTILINE),  # 数字标题 (1. 1.1 等)
            re.compile(r"^\s*[一二三四五六七八九十]\w*[章节部分]\s+.*$", re.MULTILINE),  # 中文章节
            re.compile(r"^\s*[A-Z][A-Z\s]*[：:]\s*.*$", re.MULTILINE),  # 大写字母标题
            re.compile(r"^\s*作业\d+.*$", re.MULTILINE),  # 作业标题
        ]

        lines = content.splitlines()
        current_title = None
        current_content = []

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                if current_title:
                    current_content.append(line)
                continue

            # 检查是否匹配任何标题模式
            is_header = False
            for pattern in header_patterns:
                if pattern.match(line_stripped):
                    is_header = True
                    break

            if is_header:
                # 发现新标题，保存之前的章节
                if current_title and current_content:
                    page_content = "\n".join(current_content).strip()
                    if page_content:  # 只有内容不为空时才创建Document
                        metadata = {"title": current_title.strip()}
                        sections.append(Document(page_content=page_content, metadata=metadata))

                # 开始新章节
                current_title = line_stripped
                current_content = []
            else:
                # 非标题行，添加到当前章节内容
                if current_title:
                    current_content.append(line)

        # 处理最后一个章节
        if current_title and current_content:
            page_content = "\n".join(current_content).strip()
            if page_content:
                metadata = {"title": current_title.strip()}
                sections.append(Document(page_content=page_content, metadata=metadata))

        # 如果没有找到任何标题，将整个文档作为一个章节
        if not sections and content.strip():
            metadata = {"title": "整个文档"}
            sections.append(Document(page_content=content.strip(), metadata=metadata))

        return sections

    def get_file_hash(self, file_path: str) -> str:
        with open(file_path, 'rb') as f:
            print(f"开始计算docx文档hash: {os.path.basename(file_path)}")
            return hashlib.md5(f.read()).hexdigest()

    @property
    def custom_chunking(self) -> bool:
        """已实现自定义分块"""
        return True

class XMLProcessor(DataProcessor):
    """
    专门处理UML流程模型XML的处理器
    继承自DataProcessor，实现自定义分块
    """
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理UML XML文件"""
        if not file_path.lower().endswith('.xml'):
            return False
        try:
            # 快速检查XML结构是否匹配UML模型
            tree = ET.parse(file_path)
            root = tree.getroot()

            # 检查是否包含UML模型的关键元素
            if root.find(".//phase") is not None and root.find(".//transition") is not None:
                print(f"检测到UML流程模型XML文件: {os.path.basename(file_path)}")
                return True
        except ET.ParseError:
            pass
        return False

    def load_documents(self, file_path: str) -> List[Document]:
        """
        加载UML流程模型XML文件，并将其中的阶段和转换分别转换为文档块。

        Args:
            file_path (str): XML文件的路径。

        Returns:
            List[Document]: 代表节点和关系的文档块列表。
        """
        print(f"开始处理XML文件: {os.path.basename(file_path)}")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        documents = []
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            # 提取顶层模型信息
            model_name = root.get("name", "未命名模型")
            model_domain = root.get("domain", "未知领域")

            # 处理所有图的节点
            phase_nodes = root.findall(".//phase")
            # 创建一个ID到名称的映射，方便后续处理转换关系时查找
            phase_id_to_name_map = {p.get("id"): p.get("name") for p in phase_nodes}

            for phase in phase_nodes:
                phase_id = phase.get("id")
                phase_name = phase.get("name")
                description = phase.find("description").text if phase.find("description") is not None else ""
                outcome = phase.find("outcome").text if phase.find("outcome") is not None else ""

                # 构建人类可读的文本内容
                page_content = f"阶段名称: {phase_name} (阶段ID: {phase_id}). 描述: {description}. 主要产出: {outcome}."

                # 构建丰富的元数据
                metadata = {
                    "model_name": model_name,
                    "model_domain": model_domain,
                    "element_type": "phase",  # 关键信息：这是一个节点
                    "phase_id": phase_id,
                    "phase_name": phase_name
                }
                documents.append(Document(page_content=page_content, metadata=metadata))

            # 处理所有图的边
            transition_nodes = root.findall(".//transition")

            for i, transition in enumerate(transition_nodes):
                from_id = transition.get("from")
                to_id = transition.get("to")
                condition = transition.get("condition")

                # 使用之前创建的映射表，将ID转换为可读的名称
                from_name = phase_id_to_name_map.get(from_id, f"未知阶段({from_id})")
                to_name = phase_id_to_name_map.get(to_id, f"未知阶段({to_id})")

                # 构建人类可读的文本内容
                page_content = f"流程转换: 从阶段 '{from_name}' 到 '{to_name}'. 触发条件: {condition}."

                # 构建丰富的元数据
                metadata = {
                    "model_name": model_name,
                    "model_domain": model_domain,
                    "element_type": "transition",  # 关键信息：这是一个关系
                    "from_phase_id": from_id,
                    "to_phase_id": to_id,
                    "condition": condition,
                    "transition_number": i + 1
                }
                documents.append(Document(page_content=page_content, metadata=metadata))

            print(f"模型 {os.path.basename(file_path)} 加载与解析成功！\n共生成 {len(documents)} 个文档块。")
            return documents

        except ET.ParseError as e:
            print(f"处理XML文档失败：{e}")
            raise
        except Exception as e:
            print(f"处理XML文档失败：{e}")
            raise

    def get_file_hash(self, file_path: str) -> str:
        """计算文件hash"""
        print(f"开始计算XML文档hash: {os.path.basename(file_path)}")
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()

    @property
    def custom_chunking(self) -> bool:
        """已实现自定义分块"""
        return True
class JSONProcessor(DataProcessor):
    """JSON文档处理器"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理JSON文件"""
        if file_path.lower().endswith('.json'):
            print(f"检测到JSON文件: {os.path.basename(file_path)}")
            return True
        return False

    def load_documents(self, file_path: str) -> List[Document]:
        """加载JSON文档并转换为Document对象"""
        print(f"开始处理JSON文档: {os.path.basename(file_path)}")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            documents = []

            # 判断JSON结构类型并相应处理
            if isinstance(data, dict):
                documents.extend(self._process_dict(data, file_path))
            elif isinstance(data, list):
                documents.extend(self._process_list(data, file_path))
            else:
                documents.append(self._create_document(
                    content=str(data),
                    metadata={'type': 'simple_value'}
                ))

            # 添加调试信息和文档内容验证
            print(f"成功处理JSON文档，生成 {len(documents)} 个文档片段")
            for i, doc in enumerate(documents):
                print(f"文档 {i + 1}: 类型={doc.metadata.get('type')}, 内容长度={len(doc.page_content)}")
                if doc.metadata.get('type') == 'question':
                    print(
                        f"  题目ID: {doc.metadata.get('question_id')}, 题目编号: {doc.metadata.get('question_number')}")

                # 验证文档内容完整性
                if len(doc.page_content.strip()) == 0:
                    print(f"  警告: 文档{i + 1}内容为空!")

            return documents

        except Exception as e:
            print(f"处理JSON文档失败: {e}")
            import traceback
            traceback.print_exc()
            return []

    def _process_dict(self, data: dict, file_path: str) -> List[Document]:
        """处理字典类型的JSON数据"""
        documents = []

        # 检查是否为课程练习题格式
        if 'questions' in data and isinstance(data['questions'], list):
            documents.extend(self._process_course_questions(data, file_path))
        else:
            documents.extend(self._process_generic_dict(data, file_path))

        return documents

    def _process_course_questions(self, data: dict, file_path: str) -> List[Document]:
        """处理课程练习题库格式的JSON"""
        documents = []
        # 为每个题目创建独立的详细文档
        for i, question in enumerate(data.get('questions', []), 1):
            question_doc = self._process_single_question_enhanced(question, data, file_path, i)
            documents.append(question_doc)

        return documents
    def _process_single_question_enhanced(self, question: dict, course_data: dict, file_path: str,
                                          question_number: int) -> Document:
        """处理单个题目 """
        question_id = question.get('id', 'unknown')
        question_type = question.get('type', 'unknown')
        question_text = question.get('text', '')

        # 构建题目内容
        content_parts = []
        content_parts.append(f"题目库: {os.path.splitext(os.path.basename(file_path))[0]}")
        content_parts.append(f"题目编号: 第{question_number}题")
        content_parts.append(f"题目ID: {question_id}")
        content_parts.append(f"题目类型: {question_type}")
        content_parts.append(f"难度等级: {question.get('difficulty', 'unknown')}")
        content_parts.append(f"题目内容: {question_text}")

        # 处理选择题选项
        if 'options' in question and question['options']:
            content_parts.append("\n选项:")
            for option in question['options']:
                content_parts.append(f"  {option.get('id', '')}: {option.get('text', '')}")

        # 处理答案
        if 'answer' in question:
            content_parts.append("")
            answer = question['answer']
            if isinstance(answer, list):
                content_parts.append(f"标准答案: {', '.join(map(str, answer))} ")
            else:
                content_parts.append(f"标准答案: {answer}")

        # 处理解释
        if 'explanation' in question:
            content_parts.append(f"\n答案解释: {question['explanation']}")

        # 处理编程题代码
        if question_type == 'coding' and 'solution' in question:
            solution = question['solution']
            if isinstance(solution, dict) and 'code' in solution:
                content_parts.append(f"\n参考代码:\n{solution['code']}")

        # 处理论述题指导
        if question_type == 'essay' and 'answer_guide' in question:
            guide = question['answer_guide']
            content_parts.append("\n答题指导:")
            if 'sections' in guide:
                for section in guide['sections']:
                    content_parts.append(f"  {section.get('title', '')}:")
                    for point in section.get('points', []):
                        content_parts.append(f"    - {point}")

        # 处理标签
        if 'tags' in question and question['tags']:
            content_parts.append(f"\n相关标签: {', '.join(question['tags'])}")

        content = '\n'.join(content_parts)

        # 构建元数据
        metadata = {
            'type': 'question',
            'question_id': question_id,
            'question_number': question_number,
            'question_type': question_type,
            'difficulty': question.get('difficulty', 'unknown'),
            'course_title': course_data.get('title', ''),
            'tags': question.get('tags', []),
            'has_answer': 'answer' in question,
            'has_explanation': 'explanation' in question
        }

        return self._create_document(content=content, metadata=metadata)

    def _process_generic_dict(self, data: dict, file_path: str) -> List[Document]:
        """处理通用字典格式"""
        documents = []
        main_content = self._dict_to_text(data)
        documents.append(self._create_document(
            content=main_content,
            metadata={'type': 'dict_main'}
        ))
        return documents

    def _process_list(self, data: list, file_path: str) -> List[Document]:
        """处理列表类型的JSON数据"""
        documents = []
        for i, item in enumerate(data):
            content = self._dict_to_text(item) if isinstance(item, dict) else str(item)
            documents.append(self._create_document(
                content=content,
                metadata={'type': 'list_item', 'item_index': i}
            ))
        return documents

    def _dict_to_text(self, data: dict, indent: int = 0) -> str:
        """将字典转换为可读文本"""
        lines = []
        prefix = "  " * indent
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(self._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        lines.append(f"{prefix}  [{i}]:")
                        lines.append(self._dict_to_text(item, indent + 2))
                    else:
                        lines.append(f"{prefix}  [{i}]: {item}")
            else:
                lines.append(f"{prefix}{key}: {value}")
        return '\n'.join(lines)

    def _create_document(self, content: str, metadata: dict) -> Document:
        """创建Document对象"""
        return Document(page_content=content, metadata=metadata)

    def get_file_hash(self, file_path: str) -> str:
        """计算文件hash"""
        print(f"开始计算JSON文档hash: {os.path.basename(file_path)}")
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()

    @property
    def custom_chunking(self) -> bool:
        """已实现自定义分块"""
        return True